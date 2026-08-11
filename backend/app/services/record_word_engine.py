# -*- coding: utf-8 -*-
"""受控原始记录 Word 引擎 — 填入实验数据 + 电子签名

Ported from bp-lims-reference/template_record_engine.py + record_word_engine.py
"""
from __future__ import annotations

import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(255, 0, 0)

BLANK_RE = re.compile(r"_{2,}|＿{2,}|…{2,}")
SPACE_RE = re.compile(r"\s+")
DATE_RE = re.compile(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})")
TIME_RE = re.compile(r"(\d{1,2}):(\d{1,2})(?::(\d{1,2}))?")

# ── 工具函数 ──────────────────────────────────────────────

def _clean(value: Any) -> str:
    return SPACE_RE.sub(" ", str(value or "").replace("\xa0", " ")).strip()


def _fill_date(original: str, raw: str) -> str:
    match = DATE_RE.search(str(raw or ""))
    if not match:
        return original
    year, month, day = match.groups()
    groups = list(BLANK_RE.finditer(original))
    if len(groups) < 3:
        return BLANK_RE.sub(str(raw), original, count=1)
    replacements = [year, f"{int(month):02d}", f"{int(day):02d}"]
    output = original
    for marker, replacement in zip(groups, replacements):
        start, end = marker.start(), marker.end()
        output = output[:start] + str(replacement) + " " * (end - start - len(str(replacement))) + output[end:]
    return output


def _fill_time(original: str, raw: str) -> str:
    match = TIME_RE.search(str(raw or ""))
    if not match:
        return original
    hour, minute, second = match.groups()
    replacements = [f"{int(hour):02d}", f"{int(minute):02d}"]
    if second is not None:
        replacements.append(f"{int(second):02d}")
    groups = list(BLANK_RE.finditer(original))
    output = original
    for marker, replacement in zip(groups, replacements):
        start, end = marker.start(), marker.end()
        output = output[:start] + str(replacement) + " " * (end - start - len(str(replacement))) + output[end:]
    return output


def _compose_cell_text(original: str, raw_value: Any) -> str:
    raw = str(raw_value or "").strip()
    if not original:
        return raw
    if not raw:
        return original
    if ("□" in original or "☐" in original) and ("☑" in raw or "□" in raw):
        return raw
    if "年" in original and "月" in original and "日" in original and BLANK_RE.search(original):
        return _fill_date(original, raw)
    if ":" in original and BLANK_RE.search(original) and TIME_RE.search(raw):
        return _fill_time(original, raw)
    if BLANK_RE.search(original):
        return BLANK_RE.sub(raw, original, count=1)
    return raw


def _contains_marker(text: str) -> bool:
    value = _clean(text)
    return not value or bool(BLANK_RE.search(value)) or "□" in value or "☐" in value


def _unique_row_cells(row) -> list[tuple[int, Any, str]]:
    result: list[tuple[int, Any, str]] = []
    seen = set()
    for col_index, cell in enumerate(row.cells):
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        result.append((col_index, cell, _clean(cell.text)))
    return result


def _table_headers(table) -> dict[int, str]:
    if not table.rows:
        return {}
    return {col_index: text for col_index, _, text in _unique_row_cells(table.rows[0])}


def _body_table_sections(doc: _Document) -> list[str]:
    sections: list[str] = []
    last_text = ""
    table_index = 0
    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            text = _clean(Paragraph(child, doc._body).text)
            if text:
                last_text = text
        elif child.tag == qn("w:tbl"):
            sections.append(last_text or f"表{table_index + 1}")
            table_index += 1
    return sections


def _infer_input_type(original: str) -> str:
    if "□" in original or "☐" in original:
        return "checkbox"
    if "年" in original and "月" in original and "日" in original and BLANK_RE.search(original):
        return "date"
    if len(original) > 45:
        return "textarea"
    return "text"


def _field_label(
    table_index: int, row_index: int, col_index: int,
    row_cells: list, headers: dict[int, str], template_text: str,
) -> tuple[str, str, str]:
    row_label = ""
    immediate_left = ""
    for candidate_col, _, candidate_text in row_cells:
        if candidate_col >= col_index:
            break
        if candidate_text:
            row_label = candidate_text
            immediate_left = candidate_text
    col_header = headers.get(col_index, "")

    meaningful_template = re.sub(r"[_＿]+", "", template_text).strip(" /：:；;")
    prefix = ""
    if template_text and meaningful_template and not template_text.startswith("□"):
        prefix = re.split(r"_{2,}|＿{2,}|□", template_text, maxsplit=1)[0].strip(" /：:；;")

    if not template_text and immediate_left and not _contains_marker(immediate_left):
        label = immediate_left
    elif col_header and not _contains_marker(col_header):
        label = col_header
    elif prefix:
        label = prefix
    elif row_label and not _contains_marker(row_label):
        label = row_label
    else:
        label = f"表{table_index + 1}第{row_index + 1}行第{col_index + 1}列"

    if row_index > 0 and col_header and label == col_header:
        label = f"{label}（第{row_index}条）"
    return label, row_label, col_header


# ── 模板清单 ──────────────────────────────────────────────

def template_manifest(template_path: Path | str) -> list[dict[str, Any]]:
    """解析 DOCX 模板，返回所有可填充字段的清单"""
    path = Path(template_path)
    if not path.exists():
        return []
    doc = Document(str(path))
    sections = _body_table_sections(doc)
    fields: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables):
        headers = _table_headers(table)
        seen_cells = set()
        for row_index, row in enumerate(table.rows):
            row_cells = _unique_row_cells(row)
            for col_index, cell, template_text in row_cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                if row_index == 0 and not _contains_marker(template_text):
                    continue
                if not _contains_marker(template_text):
                    continue
                label, row_label, col_header = _field_label(
                    table_index, row_index, col_index, row_cells, headers, template_text
                )
                fields.append({
                    "key": f"t{table_index}_r{row_index}_c{col_index}",
                    "section": sections[table_index] if table_index < len(sections) else f"表{table_index + 1}",
                    "table": table_index, "row": row_index, "col": col_index,
                    "label": label, "row_label": row_label, "col_header": col_header,
                    "template_text": template_text,
                    "input_type": _infer_input_type(template_text),
                    "position": f"表{table_index + 1}-R{row_index + 1}C{col_index + 1}",
                })
    return fields


# ── 单元格填入 ────────────────────────────────────────────

def _clone_rpr(source_run, target_run):
    """Copy run properties from source to target."""
    if source_run is None:
        return
    src_rpr = source_run._r.find(qn("w:rPr"))
    if src_rpr is None:
        return
    tgt_rpr = target_run._r.find(qn("w:rPr"))
    if tgt_rpr is None:
        from lxml import etree
        tgt_rpr = etree.SubElement(target_run._r, qn("w:rPr"))
    for child in list(src_rpr):
        tgt_rpr.append(deepcopy(child))


def _write_cell_text(cell, original: str, value: Any, changed: bool = False) -> None:
    text = "" if value is None else str(value)
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    source_run = next((r for p in paragraphs for r in p.runs), None)
    runs = [run for p in paragraphs for run in p.runs]
    if not runs:
        runs = [paragraph.add_run("")]
        _clone_rpr(source_run, runs[0])
    runs[0].text = text
    runs[0].font.color.rgb = RED if changed else BLACK
    for run in runs[1:]:
        run.text = ""
        run.font.color.rgb = BLACK


def fill_exact_template(
    template_path: Path | str,
    values: dict[str, Any],
    changed_keys: set[str] | None = None,
) -> _Document:
    """将 values 精确填入模板对应单元格，不增删任何表格结构"""
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"受控原始记录模板不存在：{template_path}")
    changed_keys = changed_keys or set()
    doc = Document(str(path))
    manifest_map = {field["key"]: field for field in template_manifest(path)}
    for key, value in values.items():
        field = manifest_map.get(key)
        if not field:
            continue
        table = doc.tables[field["table"]]
        row = table.rows[field["row"]]
        if field["col"] >= len(row.cells):
            continue
        cell = row.cells[field["col"]]
        _write_cell_text(cell, str(field.get("template_text", "") or ""), value, key in changed_keys)
    return doc


# ── 电子签名 ──────────────────────────────────────────────

def _signature_path(signature_dir: Path, username: str) -> Path | None:
    """查找用户签名图片"""
    if not username or not signature_dir.exists():
        return None
    for ext in (".png", ".jpg", ".jpeg"):
        candidate = signature_dir / f"{username}{ext}"
        if candidate.exists():
            return candidate
    return None


def _put_signature(cell, signature_dir: Path, username: str, date_text: str = ""):
    """在单元格中放置签名图片"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    path = _signature_path(signature_dir, username)
    if path:
        paragraph.add_run().add_picture(str(path), width=Inches(0.82))
    else:
        paragraph.add_run("【未配置签名图片】")
    if date_text:
        paragraph.add_run(f"  {str(date_text)[:10]}")


def _apply_record_signatures(
    doc: _Document, record: dict, task: dict | None,
    signature_dir: Path,
):
    """在 DOCX 中填入实验员/复核员的电子签名"""
    tester = record.get("owner") or (task.get("assignee") if task else "")
    reviewer = task.get("reviewer") if task else ""
    tester_date = record.get("tester_signed_at") or record.get("updated_at") or ""
    reviewer_date = record.get("reviewer_signed_at") or ""

    tester_tokens = ("检测人员", "实验员")
    reviewer_tokens = ("核验人员", "复核人员", "核验员", "复核员")

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        observer_columns = [i for i, text in enumerate(headers) if "观察者签字" in text]
        for row in table.rows[1:]:
            for col in observer_columns:
                if col < len(row.cells) and row.cells[col].text.strip() not in ("", "/", "不适用"):
                    _put_signature(row.cells[col], signature_dir, tester, tester_date)

        for row_index, row in enumerate(table.rows):
            cells = []
            seen = set()
            for index, cell in enumerate(row.cells):
                if cell._tc in seen:
                    continue
                seen.add(cell._tc)
                cells.append((index, cell, cell.text.strip()))
            row_text = " ".join(text for _, _, text in cells)
            exact_role_row = any(
                text.strip(" ：:（）()") in tester_tokens + reviewer_tokens
                for _, _, text in cells
            )
            signature_row = (
                any(token in row_text for token in ("签字", "签名", "日期", "年__", "年__月", "/年/月"))
                or (row_index > 0 and exact_role_row and len(cells) <= 8)
            )
            if not signature_row:
                continue
            for position, (_index, _cell, label) in enumerate(cells):
                username = ""
                signed_at = ""
                if any(token in label for token in tester_tokens):
                    username, signed_at = tester, tester_date
                elif any(token in label for token in reviewer_tokens):
                    username, signed_at = reviewer, reviewer_date
                if not username or position + 1 >= len(cells):
                    continue
                target_position = position + 1
                for later in range(position + 1, len(cells) - 1):
                    if cells[later][2].strip("：: ") in ("签字", "签名"):
                        target_position = later + 1
                        break
                _put_signature(cells[target_position][1], signature_dir, username, signed_at)

            # 合并确认格（如 R004）
            for _index, cell, text in cells:
                if "确认人" in text and "复核" in text:
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.add_run("确认：")
                    tester_path = _signature_path(signature_dir, tester)
                    if tester_path:
                        p.add_run().add_picture(str(tester_path), width=Inches(0.62))
                    p.add_run("  复核：")
                    reviewer_path = _signature_path(signature_dir, reviewer)
                    if reviewer_path:
                        p.add_run().add_picture(str(reviewer_path), width=Inches(0.62))


# ── 回退模板 ──────────────────────────────────────────────

def _fallback_docx(record: dict) -> BytesIO:
    """没有受控模板时的兜底 DOCX"""
    payload = record.get("payload", {})
    if isinstance(payload, str):
        import json
        try:
            payload = json.loads(payload)
        except Exception:
            payload = {}
    doc = Document()
    doc.add_heading(record.get("experiment") or "实验原始记录", 0)
    doc.add_paragraph("当前实验尚未配置受控原始记录模板。")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "记录值"
    template_fields = payload.get("template_fields") or payload.get("_template_fields") or {}
    if isinstance(template_fields, list):
        for tf in template_fields:
            cells = table.add_row().cells
            cells[0].text = str(tf.get("key") or tf.get("label") or "")
            cells[1].text = str(tf.get("value") or "")
    else:
        for key, value in template_fields.items():
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── 导出入口 ──────────────────────────────────────────────

# DB kind → Mapper kind
_KIND_TO_MAPPER = {
    "roughness": "rough",
    "crack": "mc_crack",
    "xray": "xray",
    "warpage": "warp",
    "cte": "cte",
    "thermal_shock": "shock",
    "bending": "bend",
    "vickers": "hv",
    "thickness": "thickness",
    "color_stability": "color",
    "fixed_denture": "fixed_denture",
    "removable_denture": "removable_denture",
    "density": "density",
    "tarnish": "tarnish",
}

# experiment_code → mapper kind 的直接映射（用于 DB 中无 kind 字段的回退）
_CODE_TO_MAPPER = {
    "I001": "rough", "I002": "mc_crack", "I003": "xray",
    "I004": "warp", "I005": "cte", "I006": "shock",
    "I007": "bend", "I008": "hv", "I009": "thickness",
    "I010": "color", "I011": "fixed_denture", "I012": "removable_denture",
    "I013": "density", "I014": "tarnish",
}


def _get_kind(record: dict, task: dict | None) -> str:
    """从 record/task 提取 mapper kind 值（多级回退）"""
    kind = record.get("kind") or (task.get("kind") if task else None) or ""
    if kind:
        mapped = _KIND_TO_MAPPER.get(kind)
        if mapped:
            return mapped
    # 通过 experiment_code 回退
    exp_code = record.get("experiment_code") or (task.get("experiment_code") if task else None) or ""
    if exp_code:
        mapped = _CODE_TO_MAPPER.get(exp_code)
        if mapped:
            return mapped
    return kind


def export_record_docx(
    record: dict,
    task: dict | None,
    template_dir: Path,
    signature_dir: Path,
) -> bytes:
    """生成填入实验数据 + 电子签名的受控 DOCX 字节流

    record: 记录字典（含 payload）
    task:   任务字典（含 reviewer, assignee 等）
    """
    payload = record.get("payload", {})
    if isinstance(payload, str):
        import json
        try:
            payload = json.loads(payload)
        except Exception:
            payload = {}

    # 确定模板文件
    template_name = record.get("record_template_file") or payload.get("_template_name") or ""
    if not template_name:
        # 尝试从实验代码推断 template_code
        experiment_code = task.get("experiment_code") if task else record.get("experiment_code", "")
        # Resolve kind → template code via the same mapping used in experiment_config
        kind = (task.get("kind") if task else None) or record.get("kind", "")
        if not kind and experiment_code:
            # Try to look up kind from experiment_methods
            kind = experiment_code  # fallback

        # Map kind / experiment_code to template code
        _KIND_MAP = {
            "rough": "R001", "mc_crack": "R004", "xray": "R005",
            "warp": "R006", "cte": "R007", "shock": "R009",
            "bend": "R010", "hv": "R011", "color": "R012", "thickness": "R013",
            "fixed_denture": "R014", "removable_denture": "R015",
            "density": "R016", "tarnish": "R017",
        }
        _CODE_TO_TEMPLATE = {
            "I001": "R001", "I002": "R004", "I003": "R005",
            "I004": "R006", "I005": "R007", "I006": "R009",
            "I007": "R010", "I008": "R011", "I009": "R013",
            "I010": "R012", "I011": "R014", "I012": "R015",
            "I013": "R016", "I014": "R017",
        }
        template_code = _KIND_MAP.get(kind) or _CODE_TO_TEMPLATE.get(experiment_code, experiment_code)

        # Search templates: try R001_*.docx, then RECORD_R001_*.docx, then SOP_R001_*.docx
        if template_dir.exists() and template_code:
            for prefix in (template_code + "_", template_code + ".", "RECORD_" + template_code, "SOP_" + template_code):
                for f in template_dir.iterdir():
                    if f.suffix != '.docx':
                        continue
                    if f.name.startswith(prefix):
                        template_name = f.name
                        break
                if template_name:
                    break

    template_path = template_dir / template_name if template_name else None
    if not template_name or not template_path or not template_path.exists():
        return _fallback_docx(record).getvalue()

    # 提取模板字段值（兼容 _template_fields 和 template_fields 两种键名）
    template_fields = payload.get("template_fields") or payload.get("_template_fields") or {}
    if isinstance(template_fields, list):
        template_fields = {tf.get("key"): tf.get("value") for tf in template_fields}

    # 始终尝试通过受控映射填充空值
    try:
        from app.services.controlled_template_mappings import apply_controlled_mapping
        kind = _get_kind(record, task)
        if kind and template_name:
            template_path_full = template_dir / template_name
            context = {
                "experiment": record.get("experiment") or (task.get("experiment") if task else ""),
                "experiment_code": record.get("experiment_code") or (task.get("experiment_code") if task else ""),
                "tester": record.get("owner", ""),
            }
            # 构建 mapper 期望的 business_record 结构（parameters + rows）
            business_record = dict(payload)
            business_record.setdefault("parameters", payload.get("_form") or {})
            business_record.setdefault("rows", payload.get("_rows") or [])
            mapped_fields = apply_controlled_mapping(
                str(template_path_full), kind, {}, context, business_record, ""
            )
            if mapped_fields:
                # 合并策略：已有存储值优先（含用户勾选数据），受控映射填补空值/占位符
                merged = dict(mapped_fields)
                existing = dict(template_fields) if template_fields else {}
                for k, v in existing.items():
                    existing_str = str(v or "")
                    # 存储值非空且不含未填充标记 → 保留用户数据
                    if existing_str.strip() and "____" not in existing_str:
                        merged[k] = v
                template_fields = merged
    except Exception as _e:
        import traceback
        traceback.print_exc()

    # 收集变化字段（用于红色标记）
    changed_keys = set()
    if int(record.get("version", 1) or 1) > 1:
        changes = record.get("changes") or []
        for item in changes:
            field_name = str(item.get("field_name", ""))
            if field_name.startswith("template_fields."):
                changed_keys.add(field_name.split("template_fields.", 1)[1])

    # 填入模板
    doc = fill_exact_template(template_path, template_fields, changed_keys)

    # 放置电子签名
    _apply_record_signatures(doc, record, task, signature_dir)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
