# -*- coding: utf-8 -*-
"""Record template engine — reads controlled Word templates and fills them with experiment data.

Adapted from bplab_v10_template_update/template_record_engine.py
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(255, 0, 0)
BLANK_RE = re.compile(r"_{2,}|＿{2,}|…{2,}")
SPACE_RE = re.compile(r"\s+")


def _clean(value: Any) -> str:
    return SPACE_RE.sub(" ", str(value or "").replace("\xa0", " ")).strip()


def _norm(value: Any) -> str:
    return re.sub(r"[\s：:：/\\、，,；;（）()\[\]【】\-—_]+", "", str(value or "")).lower()


def _contains_marker(text: str) -> bool:
    value = _clean(text)
    return not value or bool(BLANK_RE.search(value)) or "□" in value or "☐" in value


def _unique_row_cells(row) -> list[tuple[int, Any, str]]:
    result: list[tuple[int, Any, str]] = []
    seen = set()
    for col_index, cell in enumerate(row.cells):
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        result.append((col_index, cell, _clean(cell.text)))
    return result


def _table_headers(table) -> dict[int, str]:
    if not table.rows:
        return {}
    return {col_index: text for col_index, _, text in _unique_row_cells(table.rows[0])}


def _field_label(
    table_index: int,
    row_index: int,
    col_index: int,
    row_cells: list[tuple[int, Any, str]],
    headers: dict[int, str],
    template_text: str,
) -> tuple[str, str, str]:
    row_label = ""
    immediate_left = ""
    for candidate_col, _, candidate_text in row_cells:
        if candidate_col >= col_index:
            break
        if candidate_text:
            row_label = candidate_text
            immediate_left = candidate_text
    col_header = headers.get(col_index, "")

    meaningful_template = re.sub(r"[_＿]+", "", template_text).strip(" /：:；;")
    prefix = ""
    if template_text and meaningful_template and not template_text.startswith("□"):
        prefix = re.split(r"_{2,}|＿{2,}|□", template_text, maxsplit=1)[0].strip(" /：:；;")

    if not template_text and immediate_left and not _contains_marker(immediate_left):
        label = immediate_left
    elif col_header and not _contains_marker(col_header):
        label = col_header
    elif prefix:
        label = prefix
    elif row_label and not _contains_marker(row_label):
        label = row_label
    else:
        label = f"表{table_index + 1}第{row_index + 1}行第{col_index + 1}列"

    if row_index > 0 and col_header and label == col_header:
        label = f"{label}（第{row_index}条）"
    return label, row_label, col_header


def template_manifest(template_path: Path) -> list[dict[str, Any]]:
    """Extract all fillable fields from a controlled Word template."""
    doc = Document(str(template_path))
    fields: list[dict[str, Any]] = []

    section_names = []
    last_text = ""
    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = _clean(Paragraph(child, doc._body).text)
            if text:
                last_text = text
        elif child.tag == qn("w:tbl"):
            section_names.append(last_text or f"表{len(section_names) + 1}")

    for table_index, table in enumerate(doc.tables):
        section = section_names[table_index] if table_index < len(section_names) else ""
        headers = _table_headers(table)

        for row_index, row in enumerate(table.rows):
            row_cells = _unique_row_cells(row)
            for col_index, cell, cell_text in row_cells:
                value = _clean(cell_text)
                if not _contains_marker(value):
                    continue
                label, row_label, col_header = _field_label(
                    table_index, row_index, col_index, row_cells, headers, value
                )
                key = f"t{table_index}_r{row_index}_c{col_index}"
                input_type = "text"
                if "□" in value or "☐" in value:
                    input_type = "checkbox"

                fields.append({
                    "table": table_index,
                    "row": row_index,
                    "col": col_index,
                    "key": key,
                    "template_text": value,
                    "label": label,
                    "row_label": row_label,
                    "col_header": col_header,
                    "section": section,
                    "input_type": input_type,
                    "position": f"表{table_index + 1}/{row_index + 1}/{col_index + 1}",
                })
    return fields


def _select_checkbox(text: str, preferred: str) -> str:
    """Select the preferred checkbox option if present.

    Handles checkbox groups like □符合 □不符合 □合格 □不合格 etc.
    Avoids substring false-matches (e.g. '不符合' containing '符合').
    """
    normalized = _norm(preferred or "")
    if not normalized:
        return text.replace("☑", "□")
    options = [x.strip() for x in re.split(r"[□☐☑]", text)[1:] if x.strip()]
    result = text.replace("☑", "□")

    # First pass: exact / substring match excluding negation words
    matched_any = False
    for option in options:
        clean_opt = re.sub(r"[_＿…]+.*$", "", option).strip(" ：:；;，,")
        if not clean_opt:
            continue
        opt_norm = _norm(clean_opt)

        # Skip negation-matches: "不符合" should NOT match when preferred is "符合"
        if opt_norm.startswith("不") and normalized == opt_norm[1:]:
            continue  # "不符合" ≠ "符合"
        if normalized.startswith("不") and opt_norm == normalized[1:]:
            continue  # "不符合" ≠ "符合" (reverse)

        if opt_norm in normalized or normalized in opt_norm:
            result = re.sub(r"□\s*" + re.escape(clean_opt),
                            lambda m: "☑" + m.group(0)[1:], result, count=1)
            matched_any = True
        elif "符合" in opt_norm or "合格" in opt_norm or "正常" in opt_norm:
            if any(w in normalized for w in ["符合", "合格", "正常", "通过"]):
                # Double-check it's not a negation
                if not opt_norm.startswith("不"):
                    result = re.sub(r"□\s*" + re.escape(clean_opt),
                                    lambda m: "☑" + m.group(0)[1:], result, count=1)
                    matched_any = True

    # If nothing matched, try to select "其他" (other) as fallback
    if not matched_any and normalized:
        for option in options:
            clean_opt = re.sub(r"[_＿…]+.*$", "", option).strip(" ：:；;，,")
            if not clean_opt:
                continue
            opt_norm = _norm(clean_opt)
            if opt_norm in {"其他", "其它"}:
                result = re.sub(r"□\s*" + re.escape(clean_opt),
                                lambda m: "☑" + m.group(0)[1:], result, count=1)
                break

    return result


def _compose_cell_text(original: str, raw_value: Any) -> str:
    """Merge raw value into blank marker positions in the original template text."""
    if raw_value is None:
        raw_value = ""
    raw = str(raw_value)
    if not raw:
        result = re.sub(r"_{2,}|＿{2,}|…{2,}", "/", original)
        result = result.replace("☑", "□")
        return result

    original = str(original or "")
    if "□" in original or "☐" in original:
        return _select_checkbox(original, raw)

    # Date patterns like xxxx年xx月xx日
    date_match = re.match(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", raw)
    if date_match:
        year, month, day = date_match.groups()
        result = original
        result = re.sub(r"[12]\d{3}年", f"{year}年", result)
        result = re.sub(r"[01]?\d月", f"{int(month)}月", result)
        result = re.sub(r"[0-3]?\d日", f"{int(day)}日", result)
        # Fill remaining blanks with date components
        blanks = list(re.finditer(r"_{2,}|＿{2,}|…{2,}", result))
        parts = [year, month.zfill(2), day.zfill(2)]
        for i, m in enumerate(blanks):
            if i < len(parts):
                result = result[:m.start()] + parts[i] + result[m.end():]
        return result

    # Replace first blank marker with the raw value
    marker_match = BLANK_RE.search(original)
    if marker_match:
        return original[:marker_match.start()] + raw + original[marker_match.end():]
    return raw if not original else original


def _clone_format(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def _write_cell_text(cell, original: str, value: Any, changed: bool = False) -> None:
    text = "" if value is None else str(value)
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    source_run = next((r for p in paragraphs for r in p.runs), None)
    runs = [run for p in paragraphs for run in p.runs]
    if not runs:
        runs = [paragraph.add_run("")]
        _clone_format(source_run, runs[0])
    runs[0].text = text
    runs[0].font.color.rgb = RED if changed else BLACK
    for run in runs[1:]:
        run.text = ""
        run.font.color.rgb = BLACK


def _extract_form_data(payload: dict) -> dict[str, str]:
    """Extract flat key→value map from the _form section of payload."""
    result: dict[str, str] = {}
    form = payload.get("_form", {}) if isinstance(payload, dict) else {}
    if isinstance(form, dict):
        for k, v in form.items():
            result[k] = str(v) if v is not None else ""
    return result


def _extract_equipment_info(equipment_checks: list) -> dict[str, str]:
    """Extract equipment info from _equipment_checks list (keyed by normalized name)."""
    result: dict[str, str] = {}
    for eq in equipment_checks:
        if not isinstance(eq, dict):
            continue
        name = eq.get("equipment_name", "") or ""
        result.setdefault("management_no", eq.get("management_no", ""))
        result.setdefault("model", eq.get("model", ""))
        result.setdefault("serial_no", eq.get("serial_no", ""))
        result.setdefault("manufacturer", eq.get("manufacturer", ""))
        result.setdefault("calibration_time", str(eq.get("calibration_time", "")))
        result.setdefault("measuring_range", str(eq.get("measuring_range", "")))
        result.setdefault("equipment_class", str(eq.get("equipment_class", "")))
        # Also key by name-norm for lookup
        if name:
            result[_norm(name)] = eq.get("management_no", "")
    return result


def _extract_sample_nos(rows: list) -> list[str]:
    """Extract sample numbers from _rows."""
    result = []
    for row in rows:
        if isinstance(row, dict):
            sno = row.get("sample_no", "")
            if sno:
                result.append(str(sno))
    return result


def build_context_values(
    manifest: list[dict[str, Any]],
    payload: dict,
    template_name: str,
    task_info: dict | None = None,
) -> dict[str, str]:
    """Build field values from experiment payload + task info, matching template field labels.

    Data priority: _template_fields (pre-mapped keys) > _form / _rows data > task_info
    """
    task_info = task_info or {}
    values: dict[str, str] = {}

    # ── 0. Extract all payload data sources ──
    form = payload.get("_form", {}) if isinstance(payload, dict) else {}
    rows = payload.get("_rows", []) or payload.get("rows", []) or []
    rows = rows if isinstance(rows, list) else []
    equipment_checks = payload.get("_equipment_checks", []) or []
    equipment_checks = equipment_checks if isinstance(equipment_checks, list) else []
    template_fields = payload.get("_template_fields", []) or []
    template_fields = template_fields if isinstance(template_fields, list) else []

    # ── 1. First pass: apply _template_fields (direct key→value mappings from frontend) ──
    direct_map: dict[str, str] = {}
    for tf in template_fields:
        if isinstance(tf, dict):
            k = tf.get("key", "")
            v = tf.get("value", "")
            if k:
                direct_map[k] = str(v) if v is not None else ""

    # ── 2. Build lookup data from form + task_info ──
    form_data = _extract_form_data(payload)
    eq_data = _extract_equipment_info(equipment_checks)
    sample_nos = _extract_sample_nos(rows)

    # Extract key context values
    material = form_data.get("material", "") or payload.get("material", "")
    operator = task_info.get("assignee", "") or form_data.get("assignee", "")
    reviewer = task_info.get("reviewer", "") or form_data.get("reviewer", "")
    task_no = task_info.get("task_no", "") or ""
    test_date = form_data.get("test_date", "") or payload.get("test_date", "")
    equipment_name = form_data.get("equipment_name", "") or ""
    equipment_no = form_data.get("equipment_no", "") or ""
    report_no = task_info.get("report_no", "") or ""

    # ── 3. Build label→value lookup for remaining fields ──
    context: dict[str, str] = {
        "委托单位地址": form_data.get("client_address", ""),
        "委托单位": form_data.get("client_name", "") or payload.get("client_name", ""),
        "生产单位": form_data.get("production_unit", "") or payload.get("production_org_name", ""),
        "样品名称": form_data.get("sample_name", "") or payload.get("sample_name", ""),
        "材料工艺": material,
        "材料名称": material,
        "样品规格型号": form_data.get("sample_model", "") or payload.get("model", ""),
        "样品数量": str(len(sample_nos)) if sample_nos else "",
        "实验室样品编号": "、".join(sample_nos),
        "检测日期": test_date,
        "检测地点": form_data.get("location", "") or form_data.get("detection_location", ""),
        "检测依据": form_data.get("standard", "") or form_data.get("method", ""),
        "检测人员": operator,
        "记录人": operator,
        "操作人": operator,
        "核验人员": reviewer,
        "复核人员": reviewer,
        "报告编号": report_no,
        "任务编号": task_no,
        "实验任务编号": task_no,
        "原始记录编号": task_no,
        "样品编号批号": form_data.get("product_no", "") or payload.get("product_no", ""),
    }

    # ── 4. Classify fields in manifest: data rows vs fixed fields ──
    # Sample counter for row detection
    sample_field_counter: dict[tuple[int, int], int] = {}
    sample_row_set: set[tuple[int, int, int]] = set()  # (table, col) → set of row indices
    sample_col_for_table: dict[int, int] = {}  # table → sample_no column

    for field in manifest:
        header = _norm(field.get("col_header", ""))
        if header in {"样品编号", "试样编号"} and field["row"] > 0:
            key2 = (field["table"], field["col"])
            idx = sample_field_counter.get(key2, 0)
            sample_field_counter[key2] = idx + 1
            sample_col_for_table[field["table"]] = field["col"]
            sample_row_set.add((field["table"], field["col"], field["row"]))

    # ── 5. Main fill loop ──
    for field in manifest:
        key = field["key"]
        original = str(field.get("template_text", "") or "")
        header_norm = _norm(field.get("col_header", ""))
        label_norm = _norm(field.get("label", ""))
        row_norm = _norm(field.get("row_label", ""))

        # Skip if already filled by direct _template_fields
        if key in direct_map and direct_map[key]:
            values[key] = direct_map[key]
            continue

        raw = ""

        # ── 5a. Sample data row lookup ──
        if (field["table"], field["col"], field["row"]) in sample_row_set:
            # Calculate which sample row this corresponds to
            sample_rows_per_col: dict[int, list[int]] = {}
            for ft, fc, fr in sample_row_set:
                if ft == field["table"] and fc == field["col"]:
                    sample_rows_per_col.setdefault(fc, []).append(fr)
            key2 = (field["table"], field["col"])
            row_indices = sorted(sample_rows_per_col.get(field["col"], []))
            sample_idx = row_indices.index(field["row"]) if field["row"] in row_indices else -1
            if sample_idx >= 0 and sample_idx < len(sample_nos):
                if header_norm in {"样品编号", "试样编号"}:
                    raw = sample_nos[sample_idx]
            if not raw and sample_idx >= len(sample_nos):
                raw = "/"
                values[key] = _compose_cell_text(original, raw)
                continue

        # ── 5b. Equipment row matching ──
        if not raw and equipment_checks:
            for eq in equipment_checks:
                if not isinstance(eq, dict):
                    continue
                eq_name = eq.get("equipment_name", "")
                eq_name_norm = _norm(eq_name)
                if eq_name_norm and (eq_name_norm in row_norm or eq_name_norm in label_norm):
                    if header_norm in {"准确度最大允许误差测量不确定度", "准确度"}:
                        raw = str(eq.get("measuring_range", "") or eq.get("accuracy", ""))
                    elif "管理编号" in header_norm or header_norm in {"管理编号"}:
                        raw = eq.get("management_no", "")
                    elif "校准" in header_norm or "证书" in header_norm:
                        raw = str(eq.get("calibration_time", "") or eq.get("cert_no", ""))
                    elif "有效期" in header_norm:
                        raw = str(eq.get("calibration_time", ""))
                    elif "溯源" in header_norm:
                        raw = eq.get("manufacturer", "")
                    elif "型号" in header_norm:
                        raw = eq.get("model", "")
                    elif "编号" in header_norm:
                        raw = eq.get("management_no", "") or eq.get("serial_no", "")
                    else:
                        raw = eq.get("management_no", "") or eq.get("model", "")
                    break

        # ── 5c. Label-based matching (more precise: check col_header first) ──
        if not raw:
            # Col-header-first matching to avoid ambiguity
            combined = header_norm + label_norm + row_norm
            ctx_key = ""
            if "日期" in header_norm:
                date_cols = {"日期", "确认人日期", "检测日期", "测量日期"}
                raw = test_date
            elif "人员" in header_norm or "姓名" in header_norm or "签字" in header_norm:
                if "核验" in combined or "复核" in combined:
                    raw = reviewer
                elif "检测" in combined or "记录" in combined or "操作" in combined or "实验" in combined:
                    raw = operator
                elif "技术" in combined and "负责" in combined:
                    raw = ""  # leave blank for technical manager
            elif "编号" in header_norm or "管理编号" in header_norm:
                if "设备" in combined or "仪器" in combined:
                    raw = eq_data.get("management_no", "")
                elif row_norm and not _contains_marker(row_norm):
                    raw = row_norm
            else:
                # Fall back to broader label matching
                for label_key, val in context.items():
                    if val and _norm(label_key) in combined:
                        raw = val
                        ctx_key = label_key
                        break

        # ── 5d. Form data key matching for remaining fields ──
        if not raw:
            for fk, fv in form_data.items():
                if fv and _norm(fk) in combined:
                    raw = fv
                    break

        # ── 5e. Equipment data as fallback ──
        if not raw:
            for ek, ev in eq_data.items():
                if ev and _norm(ek) in combined:
                    raw = ev
                    break

        # ── 5f. Apply value ──
        if "□" in original or "☐" in original:
            values[key] = _select_checkbox(original, raw or "符合")
        else:
            values[key] = _compose_cell_text(original, raw)

    # ── 6. Fill measurement data from _rows ──
    if rows:
        _fill_measurement_rows_v2(manifest, values, rows, sample_nos)

    return values


# Mapping: normalized column header → row_data key(s)
_ROW_COLUMN_MAP: dict[str, list[str]] = {
    "试样编号": ["sample_no"],
    "样品编号": ["sample_no"],
    "样品名称": ["sample_name"],
    "原打印面方向确认": ["surface_confirm", "surface_check"],
    "测量线1raμm": ["roughness", "ra1", "测量线1"],
    "测量线2raμm": ["ra2", "测量线2"],
    "测量线3raμm": ["ra3", "测量线3"],
    "测量线1rzμm": ["rz1"],
    "测量线2rzμm": ["rz2"],
    "测量线3rzμm": ["rz3"],
    "ra平均值": ["ra_avg", "average_ra"],
    "rz平均值": ["rz_avg"],
    "单个试样平均raμm": ["avg_ra"],
    "粗糙度raμm": ["roughness"],
    "孔隙率": ["porosity"],
    "结论": ["conclusion", "result"],
    "备注": ["note"],
    "位置": ["position"],
    "表面确认": ["surface_check", "surface_confirm"],
    "表面检查": ["surface_check", "surface_confirm"],
    "判定": ["conclusion", "judgment"],
    "试样尺寸": ["sample_size", "dimension"],
    "margin_gap": ["margin_gap"],
    "limit": ["limit"],
}


def _find_row_key(header_norm: str, row_data: dict) -> str:
    """Find the value in row_data matching a column header."""
    # Direct mapping
    if header_norm in _ROW_COLUMN_MAP:
        for candidate in _ROW_COLUMN_MAP[header_norm]:
            if candidate in row_data:
                return str(row_data.get(candidate, ""))

    # Try fuzzy match: check if any row_data key normalizes to header
    for rk, rv in row_data.items():
        if _norm(rk) == header_norm:
            return str(rv) if rv is not None else ""

    # Try substring match
    for rk, rv in row_data.items():
        if header_norm in _norm(rk) or _norm(rk) in header_norm:
            return str(rv) if rv is not None else ""

    return ""


def _fill_measurement_rows_v2(
    manifest: list[dict[str, Any]],
    values: dict[str, str],
    rows: list,
    sample_nos: list[str],
) -> None:
    """Fill measurement result rows from _rows payload data.

    Strategy: for each table that has "试样编号" or "样品编号" columns,
    map each sample to a template row and fill all columns in that row.
    """
    if not rows:
        return

    # Group manifest fields by table
    by_table: dict[int, list[dict]] = {}
    for f in manifest:
        by_table.setdefault(f["table"], []).append(f)

    for ti, fields in by_table.items():
        # Find sample number column(s) in this table
        sample_cols = set()
        row_cols: dict[int, str] = {}  # col_index → header_norm
        for f in fields:
            hn = _norm(f.get("col_header", ""))
            if hn in {"试样编号", "样品编号"}:
                sample_cols.add(f["col"])
            row_cols[f["col"]] = hn

        if not sample_cols:
            continue

        # Find which rows are data rows (have sample_no fields or are after header)
        data_rows: list[int] = []
        for f in fields:
            if f["col"] in sample_cols and f["row"] > 0:
                data_rows.append(f["row"])
        data_rows = sorted(set(data_rows))

        if not data_rows:
            continue

        # Map each sample to a data row
        for ri, row_data in enumerate(rows):
            if ri >= len(data_rows):
                break
            template_row = data_rows[ri]

            if isinstance(row_data, dict):
                # Fill every column in this row with matching data
                for col_idx, header_norm in row_cols.items():
                    key = f"t{ti}_r{template_row}_c{col_idx}"
                    # Skip if already filled
                    if key in values and values[key]:
                        continue

                    if header_norm in {"试样编号", "样品编号"}:
                        val = row_data.get("sample_no", sample_nos[ri] if ri < len(sample_nos) else "")
                    else:
                        val = _find_row_key(header_norm, row_data)

                    if val:
                        original = ""
                        for f in fields:
                            if f["col"] == col_idx and f["row"] == template_row:
                                original = str(f.get("template_text", "") or "")
                                break
                        if "□" in original or "☐" in original:
                            values[key] = _select_checkbox(original, val)
                        else:
                            values[key] = _compose_cell_text(original, val)


def fill_record_template(
    template_path: Path,
    payload: dict,
    task_info: dict | None = None,
) -> bytes:
    """Fill a controlled Word record template with experiment data. Returns DOCX bytes."""
    from io import BytesIO

    if not template_path.exists():
        raise FileNotFoundError(f"模板不存在: {template_path}")

    template_name = template_path.name
    manifest = template_manifest(template_path)
    values = build_context_values(manifest, payload, template_name, task_info)

    # Also handle simple {{KEY}} paragraph placeholders
    params = payload.get("parameters", {}) if isinstance(payload, dict) else {}
    task_info = task_info or {}

    doc = Document(str(template_path))

    # Fill paragraph placeholders
    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in {**params, **task_info}.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val) if val is not None else "")

    # Fill table cells using manifest
    manifest_map = {f["key"]: f for f in manifest}
    for key, value in values.items():
        field = manifest_map.get(key)
        if not field:
            continue
        if field["table"] >= len(doc.tables):
            continue
        table = doc.tables[field["table"]]
        if field["row"] >= len(table.rows):
            continue
        row = table.rows[field["row"]]
        if field["col"] >= len(row.cells):
            continue
        cell = row.cells[field["col"]]
        original = str(field.get("template_text", "") or "")
        _write_cell_text(cell, original, value)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
