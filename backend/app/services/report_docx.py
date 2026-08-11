"""报告文档生成引擎 — 委托单/报告/发放登记/借出归还/危废处置/样品登记等 DOCX"""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor

BLACK = RGBColor(0, 0, 0)

TEMPLATE_DIR = Path(__file__).parent.parent.parent.parent / "templates"
SIGNATURE_DIR = Path(__file__).parent.parent.parent.parent / "data" / "signatures"


def _blacken(doc: Document) -> None:
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb = BLACK
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BLACK


def _save(doc: Document) -> bytes:
    _blacken(doc)
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b.read()


def _set_cell_text(cell, value: str) -> None:
    """Set cell text, preserving first paragraph."""
    paragraphs = list(cell.paragraphs)
    para = paragraphs[0] if paragraphs else cell.add_paragraph()
    runs = list(para.runs)
    if not runs:
        runs = [para.add_run("")]
    runs[0].text = "" if value is None else str(value)
    runs[0].font.color.rgb = BLACK
    for run in runs[1:]:
        run.text = ""
        run.font.color.rgb = BLACK
    for extra in paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _fill_placeholders(doc: Document, data: dict[str, str]) -> None:
    """Replace {{KEY}} placeholders in all paragraphs."""
    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val) if val is not None else "")


def _fill_table_data(table, data: list[list], header_rows: int = 1) -> None:
    """Fill table rows with data starting from header_rows."""
    while len(table.rows) < header_rows + len(data):
        table.add_row()
    for i, vals in enumerate(data, start=header_rows):
        for j, v in enumerate(vals):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = "" if v is None else str(v)
    # Clear remaining rows
    for i in range(header_rows + len(data), len(table.rows)):
        for cell in table.rows[i].cells:
            cell.text = ""


def _signature_path(username: str) -> Path | None:
    """Look up a signature image file for a username."""
    if not username:
        return None
    for ext in (".png", ".jpg", ".jpeg"):
        p = SIGNATURE_DIR / f"{username}{ext}"
        if p.exists():
            return p
    return None


def _set_signature_cell(cell, username: str, date_text: str = "", width: float = 0.92) -> None:
    """Insert signature image into a table cell."""
    paragraphs = list(cell.paragraphs)
    para = paragraphs[0] if paragraphs else cell.add_paragraph()
    para.clear()
    sig = _signature_path(username)
    if sig:
        try:
            para.add_run().add_picture(str(sig), width=Inches(width))
        except Exception:
            para.add_run("【签名图片读取失败】")
    else:
        para.add_run(f"【{username or '待签名'}】")
    if date_text:
        para.add_run(f"  {str(date_text)[:10]}")
    for extra in paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


# ═══════════════════════════════════════════════════════════════
# Report DOCX
# ═══════════════════════════════════════════════════════════════

def generate_report_docx(
    commission: dict | None,
    groups: list[dict],
    tasks: list[dict],
    records: dict[str, dict],
    report: dict,
    user_names: dict[str, str],
) -> bytes:
    """Fill FORM_REPORT.docx with all data from a completed task group."""
    template = TEMPLATE_DIR / "FORM_REPORT.docx"
    if not template.exists():
        raise FileNotFoundError(f"报告模板不存在: {template}")
    doc = Document(str(template))

    c = commission or {}
    names = "、".join(dict.fromkeys(g.get("sample_name", "") for g in groups if g.get("sample_name")))
    models = "、".join(dict.fromkeys(g.get("model", "") for g in groups if g.get("model")))

    fill_data = {
        "REPORT_NO": report.get("report_no", ""),
        "COMMISSION_NO": report.get("commission_no", ""),
        "CLIENT_NAME": c.get("client_name", ""),
        "CLIENT_ADDRESS": c.get("client_address", ""),
        "SAMPLE_NAME": names,
        "MODEL": models,
        "PRODUCT_NO": "",
        "PRODUCTION_UNIT": c.get("production_org_name", ""),
        "RECEIVE_DATE": c.get("commission_date", ""),
        "SAMPLE_CONDITION": "完好",
        "INSPECTION_CATEGORY": "委托检验",
        "REPORT_DATE": str(report.get("publish_date", "")),
        "TEST_DATE": "",
        "NOTES": report.get("notes", ""),
        "SAMPLE_STATEMENT": report.get("sample_statement", ""),
        "CONCLUSION": report.get("conclusion", ""),
        "TESTER": report.get("tester", ""),
        "VERIFIER": report.get("verifier", ""),
        "APPROVER": report.get("approver", ""),
        "EXPERIMENT": "",
    }

    # Gather experiments from tasks
    experiments = []
    test_dates = []
    for t in tasks:
        exp = t.get("experiment", "")
        if exp and exp not in experiments:
            experiments.append(exp)
        rec = records.get(t.get("task_no", ""), {})
        payload = rec.get("payload", {}) if isinstance(rec.get("payload"), dict) else {}
        params = payload.get("_form", {}) or payload.get("parameters", {})
        if params.get("test_date"):
            test_dates.append(str(params["test_date"]))

    fill_data["EXPERIMENT"] = "、".join(experiments)
    if test_dates:
        test_dates_sorted = sorted(set(test_dates))
        fill_data["TEST_DATE"] = test_dates_sorted[0] if len(test_dates_sorted) == 1 else f"{test_dates_sorted[0]}至{test_dates_sorted[-1]}"

    _fill_placeholders(doc, fill_data)

    # Equipment table (table 0 if exists)
    equipment_set = {}
    for t in tasks:
        rec = records.get(t.get("task_no", ""), {})
        payload = rec.get("payload", {}) if isinstance(rec.get("payload"), dict) else {}
        eq_checks = payload.get("_equipment_checks", [])
        for eq in eq_checks:
            key = eq.get("management_no", "")
            if key and key not in equipment_set:
                equipment_set[key] = eq

    eq_list = list(equipment_set.values())[:5]
    if eq_list and doc.tables:
        eq_table = doc.tables[0]
        eq_data = []
        for i, eq in enumerate(eq_list, 1):
            eq_data.append([
                i, eq.get("equipment_name", ""), eq.get("management_no", ""),
                eq.get("equipment_model", ""), eq.get("measuring_range", ""),
                eq.get("calibration_certificate", ""), eq.get("calibration_source", ""),
                eq.get("calibration_time", ""),
            ])
        _fill_table_data(eq_table, eq_data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Commission DOCX
# ═══════════════════════════════════════════════════════════════

def generate_commission_docx(
    commission: dict,
    groups: list[dict],
    tests: list[dict],
    receiver_name: str = "",
) -> bytes:
    """Fill FORM_COMMISSION.docx."""
    template = TEMPLATE_DIR / "FORM_COMMISSION.docx"
    if not template.exists():
        raise FileNotFoundError(f"委托单模板不存在: {template}")
    doc = Document(str(template))

    c = commission
    fill_data = {
        "CLIENT_NAME": c.get("client_name", ""),
        "CLIENT_ADDRESS": c.get("client_address", ""),
        "CONTACT": c.get("contact", ""),
        "PHONE": c.get("phone", ""),
        "COMMISSION_DATE": str(c.get("commission_date", "")),
        "DUE_DATE": str(c.get("due_date", "")),
        "PRODUCTION_UNIT": c.get("production_org_name", ""),
        "PRODUCTION_RELATION": c.get("production_relation", ""),
        "NOTES": c.get("notes", ""),
        "REPORT_MEDIUM": c.get("report_medium", "电子"),
        "CONFORMITY_JUDGMENT": c.get("conformity_judgment", ""),
        "DELIVERY_METHOD": c.get("delivery_method", ""),
    }
    _fill_placeholders(doc, fill_data)

    # Fill sample groups table
    if groups and doc.tables:
        data = []
        for i, g in enumerate(groups, 1):
            group_tests = [t.get("experiment", "") for t in tests if t.get("group_no") == g.get("group_no")]
            data.append([
                i,
                f"{g.get('sample_name', '')}（{g.get('model', '')}）",
                g.get("group_no", ""),
                c.get("production_org_name", ""),
                "、".join(group_tests),
                g.get("quantity", 1),
                g.get("notes", "") or g.get("condition_note", ""),
            ])
        _fill_table_data(doc.tables[0], data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Delivery DOCX
# ═══════════════════════════════════════════════════════════════

def generate_delivery_docx(
    report: dict,
    commission: dict | None,
    groups: list[dict],
    deliveries: list[dict],
    report_actions: list[dict] | None = None,
) -> bytes:
    """Fill FORM_REPORT_DELIVERY.docx."""
    template = TEMPLATE_DIR / "FORM_REPORT_DELIVERY.docx"
    if not template.exists():
        raise FileNotFoundError(f"报告发放模板不存在: {template}")
    doc = Document(str(template))

    c = commission or {}
    fill_data = {
        "REPORT_NO": report.get("report_no", ""),
        "COMMISSION_NO": report.get("commission_no", ""),
        "CLIENT_NAME": c.get("client_name", ""),
        "RECIPIENT": "",
        "DELIVERY_METHOD": "",
        "DELIVERY_DATE": "",
        "NOTE": "",
    }
    _fill_placeholders(doc, fill_data)

    # Fill delivery records table
    if deliveries and doc.tables:
        data = []
        for i, d in enumerate(deliveries[:7], 1):
            data.append([
                i,
                report.get("report_no", ""),
                d.get("delivery_method", ""),
                d.get("recipient", ""),
                d.get("recipient_contact", ""),
                str(d.get("delivered_at", "")),
                d.get("receipt_status", ""),
                d.get("note", ""),
            ])
        _fill_table_data(doc.tables[0], data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Sample Register DOCX
# ═══════════════════════════════════════════════════════════════

def generate_sample_register_docx(
    commission: dict,
    groups: list[dict],
    samples: list[dict],
    tests: list[dict],
    receiver_name: str = "",
) -> bytes:
    """Fill FORM_SAMPLE_REGISTER.docx."""
    template = TEMPLATE_DIR / "FORM_SAMPLE_REGISTER.docx"
    if not template.exists():
        raise FileNotFoundError(f"样品登记模板不存在: {template}")
    doc = Document(str(template))

    c = commission
    fill_data = {
        "COMMISSION_NO": c.get("commission_no", ""),
        "CLIENT_NAME": c.get("client_name", ""),
        "COMMISSION_DATE": str(c.get("commission_date", "")),
        "RECEIVER": receiver_name,
    }
    _fill_placeholders(doc, fill_data)

    gm = {g["group_no"]: g for g in groups}
    tm: dict[str, list[str]] = {}
    for t in tests:
        tm.setdefault(t["group_no"], []).append(t.get("experiment", ""))

    production_unit = c.get("production_org_name", "")

    if samples and doc.tables:
        data = []
        for s in samples:
            g = gm.get(s.get("group_no", ""), {})
            data.append([
                s.get("sample_no", ""),
                c.get("client_name", ""),
                s.get("sample_name", ""),
                s.get("model", ""),
                production_unit,
                g.get("product_no", ""),
                "、".join(tm.get(s.get("group_no", ""), [])),
                1,
                receiver_name,
                str(c.get("commission_date", "")),
                s.get("condition_note", "") or g.get("notes", ""),
            ])
        _fill_table_data(doc.tables[0], data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Sample Loan/Return DOCX
# ═══════════════════════════════════════════════════════════════

def generate_loan_return_docx(
    loans: list[dict],
    user_names: dict[str, str] | None = None,
) -> bytes:
    """Fill FORM_SAMPLE_LOAN_RETURN.docx."""
    template = TEMPLATE_DIR / "FORM_SAMPLE_LOAN_RETURN.docx"
    if not template.exists():
        raise FileNotFoundError(f"借出归还模板不存在: {template}")
    doc = Document(str(template))

    names = user_names or {}

    if loans and doc.tables:
        data = []
        for i, x in enumerate(loans, 1):
            purpose = x.get("purpose") or "、".join(json.loads(x.get("experiments", "[]")))
            data.append([
                i,
                x.get("sample_no", ""),
                names.get(x.get("borrower", ""), x.get("borrower", "")),
                str(x.get("borrowed_at", "")),
                purpose,
                str(x.get("returned_at", "")),
                names.get(x.get("returned_by", ""), x.get("returned_by", "")),
                x.get("return_note", "") or x.get("issue_note", ""),
            ])
        _fill_table_data(doc.tables[0], data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Hazardous Waste DOCX
# ═══════════════════════════════════════════════════════════════

def generate_hazardous_waste_docx(item: dict) -> bytes:
    """Fill FORM_HAZARDOUS_WASTE.docx."""
    template = TEMPLATE_DIR / "FORM_HAZARDOUS_WASTE.docx"
    if not template.exists():
        raise FileNotFoundError(f"危废处置模板不存在: {template}")
    doc = Document(str(template))

    fill_data = {
        "DISPOSAL_NO": item.get("disposal_no", ""),
        "COMMISSION_NO": item.get("commission_no", ""),
        "WASTE_NAME": item.get("waste_name", ""),
        "WASTE_TYPE": item.get("waste_type", ""),
        "QUANTITY": str(item.get("quantity", "")),
        "UNIT": item.get("unit", ""),
        "HANDLER": item.get("handler", ""),
        "CONTAINER_NO": item.get("container_no", ""),
        "OCCURRED_AT": str(item.get("occurred_at", "")),
        "DISPOSAL_METHOD": item.get("disposal_method", ""),
        "HAZARD_CATEGORY": item.get("hazard_category", ""),
        "NOTE": item.get("note", ""),
        "STATUS": item.get("status", ""),
    }
    _fill_placeholders(doc, fill_data)

    return _save(doc)


# ═══════════════════════════════════════════════════════════════
# Generic: DOCX → HTML preview helper
# ═══════════════════════════════════════════════════════════════

def docx_to_html(content: bytes, title: str = "文档预览") -> str:
    """Convert DOCX bytes to self-contained HTML for iframe preview."""
    try:
        from app.services.docx_preview import docx_review_html
        return docx_review_html(content, title)
    except ImportError:
        return f"<html><body><h2>{title}</h2><p>预览服务暂不可用</p></body></html>"
