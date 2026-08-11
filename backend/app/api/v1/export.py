"""实验数据导出 Word API — 增强版 (supports template_field_mappings + DOCX table filling)"""
from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path
from typing import Annotated

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import FileResponse, HTMLResponse, Response, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.deps import get_current_user, get_db

router = APIRouter(prefix="/export", tags=["数据导出"])


# ═══════════════════════════════════════════════════════════════
# DOCX 模板填充工具函数
# ═══════════════════════════════════════════════════════════════

def _fill_docx_paragraphs(doc, data: dict) -> None:
    """遍历所有段落，替换 {{KEY}} 占位符"""
    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val) if val is not None else "")


def _fill_docx_tables(doc, rows_data: list[dict], field_mappings: list[dict]) -> int:
    """根据 template_field_mappings 填充 DOCX 表格单元格

    Returns: 成功填充的行数
    """
    if not field_mappings or not rows_data:
        return 0

    # Group mappings by table_index
    tables_map: dict[int, list[dict]] = {}
    for m in field_mappings:
        ti = m.get("table_index", 0)
        tables_map.setdefault(ti, []).append(m)

    filled = 0
    for table_index, mappings in tables_map.items():
        if table_index >= len(doc.tables):
            continue
        table = doc.tables[table_index]

        # Sort mappings by row_index, col_index
        mappings.sort(key=lambda m: (m.get("row_index", 0), m.get("col_index", 0)))

        # Group by data source
        params_mappings = [m for m in mappings if m.get("field_source") == "params"]
        rows_mappings = [m for m in mappings if m.get("field_source") == "rows"]

        # Fill parameter fields (single-value, fill all rows with same value)
        # Find the data row that maps to params (usually row 0)
        param_data = rows_data[0] if rows_data else {}

        for m in params_mappings:
            row_idx = m.get("row_index", 0)
            col_idx = m.get("col_index", 0)
            key = m.get("field_key", "")
            transform = m.get("transform", "text")

            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    val = param_data.get("parameters", {}).get(key, param_data.get(key, ""))
                    if transform == "checkbox" and val:
                        sel = (m.get("checkbox_selection") or "").split("|")
                        cell.text = sel[0] if sel else str(val)
                    else:
                        cell.text = str(val) if val is not None else ""

        # Fill measurement rows (one row per sample)
        # mappings for row samples start after the param row
        param_row_count = max((m.get("row_index", 0) + 1) for m in params_mappings) if params_mappings else 1

        for i, sample in enumerate(rows_data):
            row_idx = param_row_count + i
            if row_idx >= len(table.rows):
                # Clone last data row if template doesn't have enough rows
                break
            for m in rows_mappings:
                col_idx = m.get("col_index", 0)
                key = m.get("field_key", "")
                if col_idx < len(table.rows[row_idx].cells):
                    val = sample.get(key, "")
                    table.rows[row_idx].cells[col_idx].text = str(val) if val is not None else ""
            filled += 1

    return filled


def _export_record_docx(template_path: Path, payload: dict, experiment_name: str,
                        task_no: str, field_mappings: list[dict] | None = None) -> Path:
    """将实验记录 payload 写入 DOCX 模板，返回输出路径"""
    from docx import Document

    doc = Document(str(template_path))

    # 1. 段落级参数填充
    params = payload.get("parameters", {}) if isinstance(payload, dict) else {}
    _fill_docx_paragraphs(doc, {
        "EXPERIMENT": experiment_name,
        "TASK_NO": task_no,
        **params,
    })

    # 2. 表格级数据填充
    if field_mappings:
        rows = payload.get("rows", []) if isinstance(payload, dict) else []
        _fill_docx_tables(doc, rows, field_mappings)
    else:
        # 无映射时追加原始数据（避免 add_heading 因为 CJK 模板样式名不同而崩溃）
        try:
            doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run("实验数据"); r.bold = True
            if isinstance(payload, dict):
                if params:
                    p2 = doc.add_paragraph()
                    r2 = p2.add_run("实验参数"); r2.bold = True
                    for key, value in params.items():
                        doc.add_paragraph(f"{key}: {value}")
                rows = payload.get("rows", [])
                if rows:
                    p3 = doc.add_paragraph()
                    r3 = p3.add_run("测量数据"); r3.bold = True
                    for i, row in enumerate(rows):
                        doc.add_paragraph(f"试样 {i+1}: {row}")
        except Exception:
            # If even this fails (e.g. style issues), just append raw data
            doc.add_page_break()
            doc.add_paragraph(str(payload))

    output_path = Path(settings.UPLOAD_DIR) / f"{task_no}_记录_v{payload.get('version', 1)}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════
# API 端点
# ═══════════════════════════════════════════════════════════════

@router.get("/record/{task_no}")
async def export_record(
    task_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    _user: Annotated[dict, Depends(get_current_user)],
):
    """导出实验原始记录为 .docx 文件（使用受控模板引擎填充）"""
    # 查询任务详细信息
    task_result = await db.execute(
        text("""
            SELECT t.experiment, t.experiment_code, t.commission_no, t.assignee, t.reviewer,
                   t.status, t.created_at
            FROM tasks t WHERE t.task_no=:t
        """),
        {"t": task_no},
    )
    task = task_result.fetchone()
    if not task:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="任务不存在")

    task_cols = task_result.keys()
    task_dict = dict(zip(task_cols, task))
    experiment_name = task_dict.get("experiment", "") or ""
    experiment_code = task_dict.get("experiment_code", "") or ""

    # 查询关联报告编号
    report_no = ""
    try:
        rp_result = await db.execute(
            text("SELECT report_no FROM reports WHERE task_no=:t ORDER BY created_at DESC LIMIT 1"),
            {"t": task_no})
        rp_row = rp_result.fetchone()
        if rp_row:
            report_no = rp_row[0] or ""
    except Exception:
        pass

    # 查询最新记录
    rec_result = await db.execute(
        text("SELECT payload, version FROM records WHERE task_no=:t ORDER BY version DESC LIMIT 1"),
        {"t": task_no},
    )
    record = rec_result.fetchone()
    if not record:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="该任务无实验记录")

    payload = json.loads(record[0]) if isinstance(record[0], str) else record[0]
    version = record[1]

    task_info = {
        "task_no": task_no,
        "experiment": experiment_name,
        "experiment_code": experiment_code,
        "report_no": report_no,
        "assignee": task_dict.get("assignee", "") or "",
        "reviewer": task_dict.get("reviewer", "") or "",
    }

    # 确定模板文件
    from app.core.encoding_rules import template_code_for_kind

    kind_result = await db.execute(
        text("SELECT kind FROM experiment_methods WHERE experiment_code=:c"),
        {"c": experiment_code},
    )
    kind_row = kind_result.fetchone()
    kind = kind_row[0] if kind_row else "generic"
    template_code = template_code_for_kind(kind)
    template_filename = f"RECORD_{template_code}.docx"
    template_path = Path(settings.TEMPLATE_DIR) / template_filename

    if not template_path.exists():
        # Fallback 1: try R-form template
        rform_candidates = sorted(Path(settings.TEMPLATE_DIR).glob(f"{template_code}_*.docx"))
        if rform_candidates:
            template_path = rform_candidates[0]
        else:
            # Fallback 2: broad search
            alt_candidates = []
            for pat in [f"RECORD_*.docx", f"R???_*.docx",
                         f"R???_*{experiment_code}*.docx", f"*{experiment_code}*.docx"]:
                alt_candidates.extend(sorted(Path(settings.TEMPLATE_DIR).glob(pat)))
            if experiment_code:
                alt_candidates.extend(sorted(Path(settings.TEMPLATE_DIR).glob(f"SOP-*_{experiment_code}*.docx")))
            if alt_candidates:
                template_path = alt_candidates[0]

    # ══ Try controlled template engine first ══
    if template_path.exists():
        try:
            from app.services.record_template_engine import fill_record_template

            docx_bytes = fill_record_template(template_path, payload, task_info)
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={task_no}_record_v{version}.docx"},
            )
        except Exception:
            pass  # Fall through to generic fallback

    # ══ Generic fallback: generate DOCX from JSON ══
    try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            doc = DocxDocument()
            # Ensure base style exists
            style = doc.styles['Normal']
            style.font.size = Pt(10)
            title = doc.add_paragraph()
            title.alignment = 1  # center
            run = title.add_run(f"实验原始记录 — {task_no}")
            run.bold = True; run.font.size = Pt(16)
            doc.add_paragraph(f"实验名称: {experiment_name or '—'}")
            doc.add_paragraph(f"实验编码: {experiment_code or '—'}")
            doc.add_paragraph("")  # spacer

            def _build_from_dict(parent_doc, data, depth=0):
                if isinstance(data, dict):
                    table = parent_doc.add_table(rows=1, cols=2, style='Table Grid')
                    table.autofit = True
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '字段'; hdr_cells[1].text = '值'
                    for p in hdr_cells[0].paragraphs: p.runs[0].bold = True if p.runs else False
                    for p in hdr_cells[1].paragraphs: p.runs[0].bold = True if p.runs else False
                    for k, v in data.items():
                        if k.startswith("_"): continue
                        if isinstance(v, (dict, list)):
                            # Section header
                            row = table.add_row()
                            row.cells[0].text = str(k)
                            row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
                            row.cells[1].text = ''
                            _build_from_dict(parent_doc, v, depth + 1)
                        else:
                            row = table.add_row()
                            row.cells[0].text = str(k)
                            row.cells[1].text = str(v) if v is not None else ''
                    parent_doc.add_paragraph("")
                elif isinstance(data, list):
                    for i, item in enumerate(data):
                        p = parent_doc.add_paragraph()
                        run = p.add_run(f"# {i + 1}")
                        run.bold = True
                        _build_from_dict(parent_doc, item, depth + 1)
                else:
                    parent_doc.add_paragraph(str(data))

            _build_from_dict(doc, payload)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(content=buf.getvalue(),
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            headers={"Content-Disposition": f"attachment; filename={task_no}_record_v{version}.docx"})
    except Exception:
        # Last resort: return JSON
        output_path = Path(settings.UPLOAD_DIR) / f"{task_no}_record_v{version}.json"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str),
                               encoding="utf-8")
        return FileResponse(
            path=str(output_path),
            filename=f"{task_no}_record_v{version}.json",
            media_type="application/json",
        )



@router.get("/report/{report_no}")
async def export_report(
    report_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    _user: Annotated[dict, Depends(get_current_user)],
):
    """导出检验报告为 .docx 文件"""
    rep_result = await db.execute(
        text("""
            SELECT r.status, r.commission_no, r.task_no, r.tester, r.verifier,
                   r.conclusion, r.notes
            FROM reports r
            WHERE r.report_no=:r
        """),
        {"r": report_no},
    )
    rep = rep_result.fetchone()
    if not rep:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="报告不存在")
    if rep[0] != "已发布":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="仅可导出已签发的报告")

    rep_cols = rep_result.keys()
    rep_dict = dict(zip(rep_cols, rep))

    # 查询任务信息获取实验名称
    task_no = rep_dict.get("task_no", "")
    experiment_name = ""
    if task_no:
        task_result = await db.execute(
            text("SELECT experiment FROM tasks WHERE task_no=:t"),
            {"t": task_no},
        )
        task_row = task_result.fetchone()
        if task_row:
            experiment_name = task_row[0] or ""

    template_path = Path(settings.TEMPLATE_DIR) / "FORM_REPORT.docx"
    if not template_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="报告模板不存在")

    try:
        from docx import Document
        doc = Document(str(template_path))

        # 查询委托信息
        comm_result = await db.execute(
            text("""
                SELECT c.commission_no, c.client_name, o.org_name AS client_org,
                       c.client_address, c.production_org_name, c.commission_date,
                       c.notes AS commission_notes, c.status AS commission_status,
                       c.method_choices
                FROM commissions c
                LEFT JOIN organizations o ON c.client_org_id = o.id
                WHERE c.commission_no=:c
            """),
            {"c": rep_dict.get("commission_no", "")},
        )
        comm = comm_result.fetchone()
        comm_info = dict(zip(comm_result.keys(), comm)) if comm else {}

        # 查询样品信息
        sample_info = {}
        try:
            sg_result = await db.execute(
                text("""
                    SELECT sample_name, model AS product_model, quantity AS sample_count,
                           material_name, product_no, condition, condition_note
                    FROM sample_groups
                    WHERE commission_no=:c AND is_void IS FALSE
                    LIMIT 1
                """),
                {"c": rep_dict.get("commission_no", "")},
            )
            sg = sg_result.fetchone()
            if sg:
                sample_info = dict(zip(sg_result.keys(), sg))
        except Exception:
            pass

        # 构建替换数据
        fill_data = {
            "报告编号：": report_no,
            "委 托 单 位：": comm_info.get("client_name", ""),
            "地       址：": comm_info.get("client_address", ""),
            "样 品 名 称：": sample_info.get("sample_name", ""),
            "型 号/规 格：": sample_info.get("product_model", ""),
            "样 品 编 号：": sample_info.get("product_no", ""),
            "产品编号/批号：": sample_info.get("product_no", ""),
            "生 产 单 位：": comm_info.get("production_org_name", ""),
            "接 收 日 期：": str(comm_info.get("commission_date", "")),
            "接 收 状 态：": sample_info.get("condition", "") or sample_info.get("condition_note", ""),
            "检 验 类 别：": rep_dict.get("report_category", ""),
            "报告发布日期：": str(rep_dict.get("publish_date", "")),
            "批 准 人": rep_dict.get("approver", ""),
            "核 验 员": rep_dict.get("verifier", ""),
            "检 测 员": rep_dict.get("tester", ""),
            "检验日期 ：": str(rep_dict.get("publish_date", "")),
            "检验结论：": rep_dict.get("conclusion", ""),
            "需说明的情况:": rep_dict.get("notes", ""),
        }

        # Fill paragraphs by matching label prefixes
        for para in doc.paragraphs:
            para_text = para.text.strip()
            for label, value in fill_data.items():
                if para_text == label or para_text.startswith(label.rstrip("：: ")):
                    # Get the first run and append value
                    if para.runs:
                        # Keep the label text, append the value
                        current = para.runs[0].text
                        val_str = str(value) if value else ""
                        if val_str:
                            if current.strip().endswith("：") or current.strip().endswith(":"):
                                para.runs[0].text = current.rstrip() + " " + val_str
                            else:
                                para.runs[0].text = current.rstrip() + val_str
                    break

        # Also fill tables with equipment/method data
        # Table 0: Equipment
        if doc.tables:
            eq_table = doc.tables[0]
            try:
                equip_result = await db.execute(
                    text("""
                        SELECT equipment_name, model, management_no,
                               calibration_cert_no, calibration_agency, calibration_due
                        FROM equipment
                        WHERE is_void IS FALSE
                        ORDER BY equipment_class
                        LIMIT 20
                    """),
                )
                equip_rows = equip_result.fetchall()
                if equip_rows and equip_result.keys():
                    eq_cols = equip_result.keys()
                    for ri, eq in enumerate(equip_rows):
                        row_idx = ri + 1
                        if row_idx >= len(eq_table.rows):
                            break
                        eq_data = dict(zip(eq_cols, eq))
                        row = eq_table.rows[row_idx]
                        mapping = [
                            eq_data.get("equipment_name", ""),
                            eq_data.get("model", ""),
                            eq_data.get("management_no", ""),
                            eq_data.get("calibration_cert_no", ""),
                            eq_data.get("calibration_agency", ""),
                            str(eq_data.get("calibration_due", "")),
                        ]
                        for ci, val in enumerate(mapping):
                            if ci < len(row.cells):
                                row.cells[ci].text = str(val) if val else ""
            except Exception:
                pass

        # Table 1: Environment (location info)
        if len(doc.tables) > 1:
            env_table = doc.tables[1]
            if len(env_table.rows) > 1 and len(env_table.rows[1].cells) >= 4:
                env_row = env_table.rows[1]
                env_row.cells[0].text = comm_info.get("client_address", "")
                # Fill temperature/humidity if available from task records
                try:
                    task_result = await db.execute(
                        text("SELECT payload FROM records WHERE task_no=:t ORDER BY version DESC LIMIT 1"),
                        {"t": task_no},
                    )
                    task_rec = task_result.fetchone()
                    if task_rec:
                        tp = json.loads(task_rec[0]) if isinstance(task_rec[0], str) else task_rec[0]
                        tf = tp.get("_form", {})
                        env_row.cells[1].text = str(tf.get("temperature_before", "")) if tf.get("temperature_before") else ""
                        env_row.cells[2].text = str(tf.get("humidity_before", "")) if tf.get("humidity_before") else ""
                except Exception:
                    pass

        # Table 2: Test method and items
        if len(doc.tables) > 2:
            method_table = doc.tables[2]
            # Fill method reference
            try:
                methods = json.loads(comm_info.get("method_choices", "[]") or "[]")
                method_names = [m.get("name", m.get("experiment", str(m))) for m in methods] if isinstance(methods, list) else []
                if method_names and method_table.rows[0].cells:
                    method_table.rows[0].cells[1].text = "; ".join(method_names[:3])
            except Exception:
                pass

            # Fill test items (row 2+) with conclusion data
            if len(method_table.rows) > 2:
                item_row = method_table.rows[2]
                if len(item_row.cells) >= 6:
                    item_row.cells[1].text = experiment_name or ""
                    item_row.cells[2].text = ""  # standard requirement
                    item_row.cells[3].text = ""  # test result
                    item_row.cells[4].text = rep_dict.get("conclusion", "")
                    item_row.cells[5].text = ""  # notes

        output_path = Path(settings.UPLOAD_DIR) / f"{report_no}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return FileResponse(
            path=str(output_path),
            filename=f"{report_no}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"导出失败: {e}")


class BatchExportRequest(BaseModel):
    task_nos: list[str] = []
    report_nos: list[str] = []


@router.post("/batch")
async def batch_export(
    body: BatchExportRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    _user: Annotated[dict, Depends(get_current_user)],
):
    """批量导出 — 将选中的任务记录和报告打包为 ZIP 下载"""
    if not body.task_nos and not body.report_nos:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="请选择要导出的任务或报告")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for tno in body.task_nos:
            try:
                rec_result = await db.execute(
                    text("SELECT payload, version FROM records WHERE task_no=:t ORDER BY version DESC LIMIT 1"),
                    {"t": tno},
                )
                rec = rec_result.fetchone()
                if rec:
                    payload = json.loads(rec[0]) if isinstance(rec[0], str) else rec[0]
                    zf.writestr(f"{tno}_记录_v{rec[1]}.json",
                                json.dumps(payload, ensure_ascii=False, indent=2, default=str))
            except Exception:
                zf.writestr(f"{tno}_error.txt", f"导出 {tno} 时出错")

        for rno in body.report_nos:
            try:
                rep_result = await db.execute(
                    text("SELECT * FROM reports WHERE report_no=:r"),
                    {"r": rno},
                )
                rep = rep_result.fetchone()
                if rep:
                    rep_dict = dict(zip(rep_result.keys(), rep))
                    zf.writestr(f"{rno}_报告.json",
                                json.dumps(rep_dict, ensure_ascii=False, indent=2, default=str))
            except Exception:
                zf.writestr(f"{rno}_error.txt", f"导出 {rno} 时出错")

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=batch_export.zip"},
    )


@router.get("/commission/{commission_no}/items")
async def list_exportable_items(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    _user: Annotated[dict, Depends(get_current_user)],
):
    """列出委托下所有可导出的任务和报告"""
    task_result = await db.execute(
        text("""
            SELECT t.task_no, t.experiment, t.status, t.assignee as tester, t.package_no
            FROM tasks t
            JOIN sample_groups sg ON t.group_no = sg.group_no
            WHERE sg.commission_no = :c
            ORDER BY t.task_no
        """),
        {"c": commission_no},
    )
    tasks = [dict(zip(task_result.keys(), r)) for r in task_result.fetchall()]

    rep_result = await db.execute(
        text("""
            SELECT report_no, status, task_no, tester
            FROM reports WHERE commission_no = :c
            ORDER BY report_no
        """),
        {"c": commission_no},
    )
    reports = [dict(zip(rep_result.keys(), r)) for r in rep_result.fetchall()]

    return {"commission_no": commission_no, "tasks": tasks, "reports": reports}


# ═══════════════════════════════════════════════════════════════
# ── 委托单预览/下载 ──
# ═══════════════════════════════════════════════════════════════

@router.get("/commission/{commission_no}/preview", response_class=HTMLResponse)
async def preview_commission(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览委托单 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    from app.services.report_docx import generate_commission_docx, docx_to_html

    c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
    c_row = c_result.fetchone()
    if not c_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="委托不存在")
    comm = dict(zip(c_result.keys(), c_row))

    g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c ORDER BY group_no"), {"c": commission_no})
    groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    t_result = await db.execute(text("SELECT * FROM tasks WHERE commission_no=:c ORDER BY task_no"), {"c": commission_no})
    tests = [dict(zip(t_result.keys(), r)) for r in t_result.fetchall()]

    receiver = comm.get("created_by", "")
    try:
        docx_bytes = generate_commission_docx(comm, groups, tests, receiver)
        html = docx_to_html(docx_bytes, f"委托单 — {commission_no}")
        return HTMLResponse(content=html)
    except FileNotFoundError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=str(e))


@router.get("/commission/{commission_no}/export")
async def download_commission(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """下载委托单 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权下载")

    from app.services.report_docx import generate_commission_docx

    c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
    c_row = c_result.fetchone()
    if not c_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="委托不存在")
    comm = dict(zip(c_result.keys(), c_row))

    g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c ORDER BY group_no"), {"c": commission_no})
    groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    t_result = await db.execute(text("SELECT * FROM tasks WHERE commission_no=:c ORDER BY task_no"), {"c": commission_no})
    tests = [dict(zip(t_result.keys(), r)) for r in t_result.fetchall()]

    receiver = comm.get("created_by", "")
    docx_bytes = generate_commission_docx(comm, groups, tests, receiver)
    return Response(content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={commission_no}_commission.docx"})


# ═══════════════════════════════════════════════════════════════
# ── 样品登记表预览/下载 ──
# ═══════════════════════════════════════════════════════════════

@router.get("/sample-register/{commission_no}/preview", response_class=HTMLResponse)
async def preview_sample_register(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览样品登记表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    from app.services.report_docx import generate_sample_register_docx, docx_to_html

    c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
    c_row = c_result.fetchone()
    if not c_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="委托不存在")
    comm = dict(zip(c_result.keys(), c_row))

    g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c ORDER BY group_no"), {"c": commission_no})
    groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    s_result = await db.execute(text("SELECT * FROM samples WHERE commission_no=:c ORDER BY sample_no"), {"c": commission_no})
    samples = [dict(zip(s_result.keys(), r)) for r in s_result.fetchall()]

    t_result = await db.execute(text("SELECT * FROM tasks WHERE commission_no=:c ORDER BY task_no"), {"c": commission_no})
    tests = [dict(zip(t_result.keys(), r)) for r in t_result.fetchall()]

    receiver = comm.get("created_by", "")
    try:
        docx_bytes = generate_sample_register_docx(comm, groups, samples, tests, receiver)
        html = docx_to_html(docx_bytes, f"样品登记表 — {commission_no}")
        return HTMLResponse(content=html)
    except FileNotFoundError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=str(e))


@router.get("/sample-register/{commission_no}/export")
async def download_sample_register(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """下载样品登记表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权下载")

    from app.services.report_docx import generate_sample_register_docx

    c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
    c_row = c_result.fetchone()
    if not c_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="委托不存在")
    comm = dict(zip(c_result.keys(), c_row))

    g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c ORDER BY group_no"), {"c": commission_no})
    groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    s_result = await db.execute(text("SELECT * FROM samples WHERE commission_no=:c ORDER BY sample_no"), {"c": commission_no})
    samples = [dict(zip(s_result.keys(), r)) for r in s_result.fetchall()]

    t_result = await db.execute(text("SELECT * FROM tasks WHERE commission_no=:c ORDER BY task_no"), {"c": commission_no})
    tests = [dict(zip(t_result.keys(), r)) for r in t_result.fetchall()]

    receiver = comm.get("created_by", "")
    docx_bytes = generate_sample_register_docx(comm, groups, samples, tests, receiver)
    return Response(content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={commission_no}_sample_register.docx"})


# ═══════════════════════════════════════════════════════════════
# ── 借出归还表预览/下载 ──
# ═══════════════════════════════════════════════════════════════

@router.get("/loan-return/{commission_no}/preview", response_class=HTMLResponse)
async def preview_loan_return(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览借出归还表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    from app.services.report_docx import generate_loan_return_docx, docx_to_html

    lr_result = await db.execute(
        text("""
            SELECT pl.* FROM package_loans pl
            JOIN samples s ON pl.sample_no = s.sample_no
            WHERE s.commission_no=:c ORDER BY pl.borrowed_at
        """),
        {"c": commission_no},
    )
    loans = [dict(zip(lr_result.keys(), r)) for r in lr_result.fetchall()]
    if not loans:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该委托无借出归还记录")

    user_names = {}
    u_result = await db.execute(text("SELECT username, display_name FROM users"))
    for u_row in u_result.fetchall():
        user_names[u_row[0]] = u_row[1] or u_row[0]

    try:
        docx_bytes = generate_loan_return_docx(loans, user_names)
        html = docx_to_html(docx_bytes, f"借出归还表 — {commission_no}")
        return HTMLResponse(content=html)
    except FileNotFoundError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=str(e))


@router.get("/loan-return/{commission_no}/export")
async def download_loan_return(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """下载借出归还表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权下载")

    from app.services.report_docx import generate_loan_return_docx

    lr_result = await db.execute(
        text("""
            SELECT pl.* FROM package_loans pl
            JOIN samples s ON pl.sample_no = s.sample_no
            WHERE s.commission_no=:c ORDER BY pl.borrowed_at
        """),
        {"c": commission_no},
    )
    loans = [dict(zip(lr_result.keys(), r)) for r in lr_result.fetchall()]
    if not loans:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该委托无借出归还记录")

    user_names = {}
    u_result = await db.execute(text("SELECT username, display_name FROM users"))
    for u_row in u_result.fetchall():
        user_names[u_row[0]] = u_row[1] or u_row[0]

    docx_bytes = generate_loan_return_docx(loans, user_names)
    return Response(content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={commission_no}_loan_return.docx"})


# ═══════════════════════════════════════════════════════════════
# ── 危废处置表预览/下载 ──
# ═══════════════════════════════════════════════════════════════

@router.get("/hazardous-waste/{commission_no}/preview", response_class=HTMLResponse)
async def preview_hazardous_waste(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览危废处置表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    from app.services.report_docx import generate_hazardous_waste_docx, docx_to_html

    hw_result = await db.execute(
        text("SELECT * FROM hazardous_waste_records WHERE commission_no=:c ORDER BY created_at DESC LIMIT 1"),
        {"c": commission_no},
    )
    hw_row = hw_result.fetchone()
    if not hw_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该委托无危废处置记录")
    item = dict(zip(hw_result.keys(), hw_row))

    try:
        docx_bytes = generate_hazardous_waste_docx(item)
        html = docx_to_html(docx_bytes, f"危废处置表 — {commission_no}")
        return HTMLResponse(content=html)
    except FileNotFoundError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=str(e))


@router.get("/hazardous-waste/{commission_no}/export")
async def download_hazardous_waste(
    commission_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """下载危废处置表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权下载")

    from app.services.report_docx import generate_hazardous_waste_docx

    hw_result = await db.execute(
        text("SELECT * FROM hazardous_waste_records WHERE commission_no=:c ORDER BY created_at DESC LIMIT 1"),
        {"c": commission_no},
    )
    hw_row = hw_result.fetchone()
    if not hw_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该委托无危废处置记录")
    item = dict(zip(hw_result.keys(), hw_row))

    docx_bytes = generate_hazardous_waste_docx(item)
    return Response(content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={commission_no}_hazardous_waste.docx"})


# ═══════════════════════════════════════════════════════════════
# ── 报告发放登记表预览/下载 ──
# ═══════════════════════════════════════════════════════════════

@router.get("/report-delivery/{report_no}/preview", response_class=HTMLResponse)
async def preview_report_delivery(
    report_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览报告发放登记表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    from app.services.report_docx import generate_delivery_docx, docx_to_html

    rep_result = await db.execute(text("SELECT * FROM reports WHERE report_no=:r"), {"r": report_no})
    rep_row = rep_result.fetchone()
    if not rep_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="报告不存在")
    report = dict(zip(rep_result.keys(), rep_row))

    del_result = await db.execute(
        text("SELECT * FROM report_deliveries WHERE report_no=:r ORDER BY delivered_at"),
        {"r": report_no},
    )
    deliveries = [dict(zip(del_result.keys(), r)) for r in del_result.fetchall()]
    if not deliveries:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该报告无发放记录")

    comm = {}
    commission_no = report.get("commission_no", "")
    if commission_no:
        c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
        c_row = c_result.fetchone()
        if c_row:
            comm = dict(zip(c_result.keys(), c_row))

    groups = []
    if commission_no:
        g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c"), {"c": commission_no})
        groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    act_result = await db.execute(
        text("SELECT * FROM report_actions WHERE report_no=:r ORDER BY created_at"),
        {"r": report_no},
    )
    actions = [dict(zip(act_result.keys(), r)) for r in act_result.fetchall()]

    try:
        docx_bytes = generate_delivery_docx(report, comm, groups, deliveries, actions)
        html = docx_to_html(docx_bytes, f"报告发放登记表 — {report_no}")
        return HTMLResponse(content=html)
    except FileNotFoundError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=str(e))


@router.get("/report-delivery/{report_no}/export")
async def download_report_delivery(
    report_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """下载报告发放登记表 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权下载")

    from app.services.report_docx import generate_delivery_docx

    rep_result = await db.execute(text("SELECT * FROM reports WHERE report_no=:r"), {"r": report_no})
    rep_row = rep_result.fetchone()
    if not rep_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="报告不存在")
    report = dict(zip(rep_result.keys(), rep_row))

    del_result = await db.execute(
        text("SELECT * FROM report_deliveries WHERE report_no=:r ORDER BY delivered_at"),
        {"r": report_no},
    )
    deliveries = [dict(zip(del_result.keys(), r)) for r in del_result.fetchall()]
    if not deliveries:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该报告无发放记录")

    comm = {}
    commission_no = report.get("commission_no", "")
    if commission_no:
        c_result = await db.execute(text("SELECT * FROM commissions WHERE commission_no=:c"), {"c": commission_no})
        c_row = c_result.fetchone()
        if c_row:
            comm = dict(zip(c_result.keys(), c_row))

    groups = []
    if commission_no:
        g_result = await db.execute(text("SELECT * FROM sample_groups WHERE commission_no=:c"), {"c": commission_no})
        groups = [dict(zip(g_result.keys(), r)) for r in g_result.fetchall()]

    act_result = await db.execute(
        text("SELECT * FROM report_actions WHERE report_no=:r ORDER BY created_at"),
        {"r": report_no},
    )
    actions = [dict(zip(act_result.keys(), r)) for r in act_result.fetchall()]

    docx_bytes = generate_delivery_docx(report, comm, groups, deliveries, actions)
    safe_name = f"{report_no}_delivery.docx"
    encoded_name = safe_name.encode("ascii", errors="ignore").decode("ascii")
    return Response(content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename=\"{encoded_name}\""})


# ═══════════════════════════════════════════════════════════════
# ── 通用 DOCX 预览 endpoint (for records) ──
# ═══════════════════════════════════════════════════════════════

@router.get("/record/{task_no}/preview", response_class=HTMLResponse)
async def preview_record_export(
    task_no: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    user: Annotated[dict, Depends(get_current_user)],
):
    """在线预览实验记录 DOCX"""
    role = user.get("role", "")
    if role not in ("质量负责人", "管理员", "复核员", "样品管理员", "实验员"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无权预览")

    rec_result = await db.execute(
        text("SELECT * FROM records WHERE task_no=:t ORDER BY version DESC LIMIT 1"),
        {"t": task_no},
    )
    rec_row = rec_result.fetchone()
    if not rec_row:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该任务无记录")
    record = dict(zip(rec_result.keys(), rec_row))

    t_result = await db.execute(text("SELECT * FROM tasks WHERE task_no=:t"), {"t": task_no})
    t_row = t_result.fetchone()
    task = dict(zip(t_result.keys(), t_row)) if t_row else None

    try:
        from app.services.record_word_engine import export_record_docx
        from app.services.docx_preview import docx_review_html
        docx_bytes = export_record_docx(record, task, template_dir=settings.TEMPLATE_DIR, signature_dir=settings.SIGNATURE_DIR)
        title = f"{record.get('experiment','原始记录')} — {task_no}"
        html = docx_review_html(docx_bytes, title)
        return HTMLResponse(content=html)
    except ImportError as e:
        raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail=f"预览服务不可用：{e}")
